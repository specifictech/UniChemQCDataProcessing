from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn

from docx2pdf import convert

import pandas as pd
from typing import Any
from datetime import datetime

import constants as const
from classes.lot import Lot

class Report:
    def __init__(self, lot):
        self.lot = lot
        self.date_generated = datetime.now().strftime("%Y-%m-%d")
        
        self.export_lot_to_docx(self.lot, output_path=f"./output/lot_report_{self.lot.lot_num}.docx")

    def export_lot_to_docx(self, lot: Lot, output_path: str):
        """
        Export a Lot object's attributes and nested InkChannel summaries to a Word document.

        Parameters:
            - lot : A Lot object containing the data to be reported.
        """
        doc = Document()

        self.set_global_page_layout(doc, doc.sections[0])
        self.set_global_styles(doc)
        
        # Page 1: Lot information & summary, report generation info
        self.add_title(doc)
        self.add_lot_info_and_test_summary(doc, lot)
        self.add_whitespace(doc, space_before_pts=12, space_after_pts=12) 
        self.add_lot_summary_result(doc, lot)
        self.add_whitespace(doc, space_before_pts=12, space_after_pts=12)
        self.add_report_generation_info(doc, lot)
        
        # Subsequent pages: spot ID + channel details
        for spot_id, channels in lot.ink_channels.items():
            self.add_whitespace(doc, page_break=True)
            self.add_heading_generic(doc, f"{spot_id} Analysis", 
                                     level=1, 
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self.add_heading_generic(doc, f"Lot: {lot.lot_num}",
                                     level=2,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     bold=False)
            self.add_whitespace(doc, space_before_pts=8, space_after_pts=8)
            
            for i in range(len(channels)):
                if i > 0: # Add page break before all but first
                    self.add_whitespace(doc, page_break=True)
                self.add_ink_channel_summary(doc, channels[i])
        
        doc.save(output_path)
        convert(output_path) # DOCX to PDF & save

    def add_two_col_row(
        self,
        container,
        left_text: str,
        right_text: str,
        width=None,
        font_size_pt: int = 10,
        embedded: bool = False,
        bold_when_shaded: bool = True,
        left_heading_level = None,
        left_color = None,
        left_space_before_pt = None,
        left_space_after_pt = None,
    ):
        """
        Create a 1x2 table and write left/right texts. Optionally render the left text
        using a Word Heading style (Heading 1..Heading 9) via `left_heading_level`.

        Parameters
        ----------
        container : Document or _Cell
            Where the table is inserted.
        left_text, right_text : str
            The texts to display.
        width : docx.shared.Length | None
            Table width. Ignored when embedded=True (python-docx constraint).
        font_size_pt : int
            Base font size for normal (non-heading) runs.
        embedded : bool
            If True, create the table without width for reliable embedding in a cell.
        shade_hex : str | None
            Optional shading (hex) for the right cell’s run.
        bold_when_shaded : bool
            If True and shade_hex is set, make the right run bold.
        left_heading_level : int | None
            If set (e.g., 2), the left text uses style "Heading {level}".
        left_color : RGBColor | None
            Optional text color for the left text (works for both heading and normal run).
        left_space_before_pt, left_space_after_pt : float | None
            Optional paragraph spacing adjustments when using heading style,
            useful inside table cells to avoid extra gaps.
        """
        # Create a 1x2 table
        if not embedded:
            tbl = container.add_table(rows=1, cols=2, width=width)
        else:
            tbl = container.add_table(rows=1, cols=2)
        tbl.style = "Normal Table"

        # --- Left cell ---
        left_cell = tbl.rows[0].cells[0]
        p_left = left_cell.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_left.clear()

        if left_heading_level is not None:
            # Apply Heading style to the paragraph
            try:
                p_left.style = f"Heading {left_heading_level}"
            except KeyError:
                # Fallback: if style isn't present, keep default paragraph style
                # (You can define Heading styles in set_global_styles if needed.)
                pass

            # Write a single run with the text
            r_left = p_left.add_run(left_text)
            if left_color is not None:
                r_left.font.color.rgb = left_color

            # Optional spacing to avoid extra gap within a table cell
            pf = p_left.paragraph_format
            if left_space_before_pt is not None:
                pf.space_before = Pt(left_space_before_pt)
            if left_space_after_pt is not None:
                pf.space_after = Pt(left_space_after_pt)
        else:
            # Normal run styling
            r_left = p_left.add_run(left_text)
            r_left.font.size = Pt(font_size_pt)
            r_left.font.bold = False
            if left_color is not None:
                r_left.font.color.rgb = left_color

        # --- Right cell ---
        right_cell = tbl.rows[0].cells[1]
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.clear()

        r_right = p_right.add_run(right_text)
        r_right.font.size = Pt(font_size_pt)
        r_right.font.bold = False

        # Optional shading on the right cell
        if bold_when_shaded:
            r_right.font.bold = True

        return tbl  # optional: return table if caller wants to further tweak

    def printable_width_from_section(self, section):
        """
        Compute the printable text width from a python-docx Section:
        page_width - left_margin - right_margin

        Returns a docx.shared.Length instance suitable for table width.
        """
        return section.page_width - section.left_margin - section.right_margin
    
    def set_global_styles(self, doc):
        h1_style = doc.styles['Heading 1']
        h1_style.font.size = Pt(16)
        h1_style.font.color.rgb = RGBColor(0, 0, 0)
        h1_style.font.bold = True
        h1_pf = h1_style.paragraph_format
        h1_pf.space_before = Pt(0)
        h1_pf.space_after = Pt(0)
        
        h2_style = doc.styles['Heading 2']
        h2_style.font.size = Pt(14)
        h2_style.font.color.rgb = RGBColor(0, 0, 0)
        h2_pf = h2_style.paragraph_format
        h2_pf.space_before = Pt(0)
        h2_pf.space_after = Pt(0)
            
    def set_global_page_layout(self, doc, section):
        section = doc.sections[-1]
        section.orientation = const.PAGE_ORIENTATION
        section.left_margin = Inches(const.LEFT_MARGIN_INCHES)
        section.right_margin = Inches(const.RIGHT_MARGIN_INCHES)
        self.set_header(section=section, software_version=self.lot.software_version)
        self.set_footer(section=section)
        
    def set_header(self, section, software_version: str):
        """
        Set the top header for a section with software_version (left)
        and date_generated (right). Applies to all pages that use this section.
        """
        header = section.header
        table_width = Inches(const.PAGE_WIDTH_INCHES)
        self.add_two_col_row(
            container=header,
            left_text=f"Software Version: {software_version}",
            right_text=f"Last Updated: {self.date_generated}",
            width=table_width,
            font_size_pt=10,
            bold_when_shaded=False
        )
        
    def set_footer(self, section):
        """Set the bottom footer for each page, displaying the confidential message."""
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(const.PDF_FOOTER)
        run.font.size = Pt(10)
        run.font.bold = False
    
    def add_title(self, doc):
        """Add the main title to the document."""
        title_style = doc.styles['Title']
        
        # Reduce space after
        pf = title_style.paragraph_format
        pf.space_after = Pt(0)
        
        # Remove default line from Title style
        pPr = title_style._element.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)
        
        # Modify font properties
        title_font = title_style.font
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        h = doc.add_heading(f"{const.PDF_TITLE}", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_lot_info_and_test_summary(self, doc, lot: Lot):
        lot_info = {
            "Lot Number:": lot.lot_num,
            "Analysis Date:": lot.analysis_date,
        }
        test_sum = {
            "Total Channels Analyzed:": str(int(lot.chan_passed) + int(lot.chan_failed)),
            "Channels Passed:": lot.chan_passed,
            "Channels Failed:": lot.chan_failed,
            "Final Lot Disposition:": lot.final_lot_disposition,
        }
        
        # Outer container for both Lot Info and Test Summary tables
        container = doc.add_table(rows=1, cols=2)
        container.style = "Normal Table"
        left_cell, right_cell = container.rows[0].cells
        left_cell.width = Inches(const.PAGE_WIDTH_INCHES/2)
        right_cell.width = Inches(const.PAGE_WIDTH_INCHES/2)
        self.kv_table_into(left_cell, lot_info, style=const.TABLE_GRID, title="Lot Information")
        self.kv_table_into(
            right_cell, 
            test_sum, 
            style=const.TABLE_GRID,
            title="Test Summary",
            left_col_width_inches=5.0
        )
        
    def add_lot_summary_result(self, doc, lot: Lot):
        df = lot.lot_summary
        self.df_into(doc, df, title="Lot Summary Result")
        
    def add_report_generation_info(self, doc, lot: Lot):
        report_gen = {
            "Operator:": lot.operator,
            "Timestamp:": lot.timestamp
        }
        self.kv_table_into(doc, report_gen, style=const.TABLE_GRID, title="Report Generation")
        
        self.add_whitespace(doc, space_before_pts=12, space_after_pts=12)
        
        self.add_heading_generic(doc, "Manual Verification", level=2)
        
        container = doc.add_table(rows=1, cols=2)
        container.style = const.NORMAL_TABLE
        
        left_cell, right_cell = container.rows[0].cells
        
        self.set_cell_margins(
            left_cell,
            left=216,
            top=const.CELL_VERTICAL_PADDING_TWIPS,
            bottom=const.CELL_VERTICAL_PADDING_TWIPS
        )
        
        p_left = left_cell.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_left.clear()
        r_left = p_left.add_run("Operator Initials:")
        r_left.font.size = Pt(13)
        r_left.font.bold = False
        r_left.add_break(WD_BREAK.LINE)
        r_left.add_break(WD_BREAK.LINE)
        p_left.add_run("____________________________________________________").font.size = Pt(13)
        
        self.set_cell_margins(
            right_cell,
            right=216,
            top=const.CELL_VERTICAL_PADDING_TWIPS,
            bottom=const.CELL_VERTICAL_PADDING_TWIPS
        )
        
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_right.clear()
        r_right = p_right.add_run("Date:")
        r_right.font.size = Pt(13)
        r_right.font.bold = False
        r_right.add_break(WD_BREAK.LINE)
        r_right.add_break(WD_BREAK.LINE)
        p_right.add_run("____________________________________________________").font.size = Pt(13)
        
    def add_ink_channel_summary(self, doc, ic):
        container = doc.add_table(rows=1, cols=1)
        container.style = const.NORMAL_TABLE
        
        host_cell = container.rows[0].cells[0]
        
        for p in list(host_cell.paragraphs):
            # Empty if no text and all runs are blank
            if not p.text.strip() and not any(r.text.strip() for r in p.runs):
                p._element.getparent().remove(p._element)
                # remove only the leading empty one
                break
            else:
                # first paragraph has content; nothing to remove
                break
        
        if 'G' in ic.channel:
            channel = 'Green'
        else: # 'R' in ic.channel
            channel = 'Red'
            
        self.add_two_col_row(
            container=host_cell,
            left_text=f"{channel} Channel Intensity",
            right_text=f"{ic.result}",
            font_size_pt=14,
            embedded = True,
            left_heading_level=2,
            left_color=RGBColor(0, 0, 0),
            left_space_before_pt=0,
            left_space_after_pt=0
        )
        
        # 🔧 Remove the auto-inserted empty paragraph that follows a table in a cell
        self._remove_trailing_empty_paragraph_in_cell(host_cell)
            
        bullet_text = f"{ic.analyte} Sensors: {ic.passing_sensors} passed, {ic.failing_sensors} failed (Total: {ic.passing_sensors + ic.failing_sensors})"
        p_bullet = host_cell.add_paragraph(bullet_text, style="List Bullet")
        p_bullet.runs[0].font.size = Pt(13)
        p_bullet.space_before = Pt(0)
        p_bullet.space_after = Pt(0)
        
        self.add_whitespace(host_cell, space_before_pts=5, space_after_pts=5)
        
        self.df_into(host_cell, 
                     ic.results_summary, 
                     title="  Results Summary", # Add spacing to align with prior heading
                     font_size=Pt(13),
                     row_space_before_pts = 2,
                     row_space_after_pts = 2,
                     row_line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE,
                     fourth_col_width_inches=0.7)
        
    def _remove_trailing_empty_paragraph_in_cell(self, cell):
        """
        Remove the last paragraph in a table cell if it is truly empty.
        Safe to call multiple times.
        """
        # python-docx keeps a live view; iterate over a copy
        for p in list(cell.paragraphs)[::-1]:  # start from the end
            # Empty if no text and all runs are blank
            if not p.text.strip() and not any(r.text.strip() for r in p.runs):
                p._element.getparent().remove(p._element)
                break  # remove only the trailing empty one
            else:
                break  # the last paragraph has content; stop
        
    def kv_table_into(
        self,
        container,
        kv: dict,
        style: str,
        title: str,
        left_col_width_inches: float = None, # type: ignore
        blank_only: bool = False
    ):  # type: ignore
        """Render a key-value table into a Document or a _Cell."""
        if title:
            self.add_heading_generic(container, title, level=2)

        # Create the table (start with 1 headerless row if you prefer)
        table = container.add_table(rows=0, cols=2)
        table.style = style

        # Turn off autofit so explicit widths are respected
        try:
            table.autofit = False
        except AttributeError:
            table.allow_autofit = False

        # Apply horizontal-only borders (your existing helper)
        self.set_table_horizontal_borders(table, blank_only)

        left_width = Inches(left_col_width_inches) if left_col_width_inches is not None else None

        for k, v in kv.items():
            row_cells = table.add_row().cells

            # >>> Ensure the left column width is set on **this** new row
            if left_width is not None:
                row_cells[0].width = left_width

            # Clear zebra shading if any (your existing code)
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    tcPr.remove(shd)

            # Left cell (key)
            p_left = row_cells[0].paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.clear()
            run_left = p_left.add_run(str(k))
            run_left.font.bold = False
            run_left.font.size = Pt(13)

            self.set_cell_margins(
                row_cells[0],
                left=216,  # ~0.15"
                top=const.CELL_VERTICAL_PADDING_TWIPS,
                bottom=const.CELL_VERTICAL_PADDING_TWIPS
            )

            # Right cell (value)
            p_right = row_cells[1].paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.clear()
            run_right = p_right.add_run(str(v))
            run_right.font.bold = False
            run_right.font.size = Pt(13)

            self.set_cell_margins(
                row_cells[1],
                right=216,  # ~0.15"
                top=const.CELL_VERTICAL_PADDING_TWIPS,
                bottom=const.CELL_VERTICAL_PADDING_TWIPS
            )
            
            cell_text = str(v)
            text_upper = cell_text.strip().upper()
            if text_upper == "PASS" or text_upper == "FAIL":
                run_right.font.bold = True

    def df_into(
        self,
        container,
        df: pd.DataFrame,
        title: str,
        font_size: Pt = Pt(13),
        row_space_before_pts = None,
        row_space_after_pts = None,
        row_line_spacing_rule = None,
        row_line_spacing = None,
        fourth_col_width_inches = None  # manual width for the 4th column
):
        # Title
        self.add_heading_generic(container, title, level=2)

        # No data case
        if df is None or df.empty:
            p = container.add_paragraph("No data.")
            p.runs[0].italic = True
            return

        # Create table with the correct number of columns
        table = container.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        # Keep only horizontal borders
        self.set_table_horizontal_borders(table)

        # Disable autofit so explicit widths are respected
        try:
            table.autofit = False
        except AttributeError:
            table.allow_autofit = False  # for older python-docx builds

        # --- Header row ---
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(df.columns):
            self.set_cell_text(hdr_cells[j], str(col), bold=True, size=font_size)
            # vertical padding
            self.set_cell_margins(
                hdr_cells[j],
                top=const.CELL_VERTICAL_PADDING_TWIPS,
                bottom=const.CELL_VERTICAL_PADDING_TWIPS
            )

        # If user asked to set width for the 4th column, apply to header cell now
        if fourth_col_width_inches is not None and len(df.columns) >= 4:
            hdr_cells[3].width = Inches(fourth_col_width_inches)

        # --- Body rows ---
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for j, col in enumerate(df.columns):
                val = row[col]
                if isinstance(val, list):
                    cell_text = ", ".join(map(str, val))
                else:
                    cell_text = "" if pd.isna(val) else str(val)

                # Clear & write a single run
                cells[j].text = ""
                p = cells[j].paragraphs[0]
                p.clear()
                r = p.add_run(cell_text)
                r.font.bold = False
                r.font.size = font_size

                pf = p.paragraph_format
                if row_space_before_pts is not None:
                    pf.space_before = Pt(row_space_before_pts)
                if row_space_after_pts is not None:
                    pf.space_after = Pt(row_space_after_pts)
                    
                # Line spacing rule + value
                if row_line_spacing_rule is not None:
                    pf.line_spacing_rule = row_line_spacing_rule
                    if row_line_spacing is not None:
                        pf.line_spacing = row_line_spacing

                # Conditional bold for PASS/FAIL
                text_upper = cell_text.strip().upper()
                if text_upper == "PASS" or text_upper == "FAIL":
                    r.font.bold = True
                
                # vertical padding
                self.set_cell_margins(
                    cells[j],
                    top=const.CELL_VERTICAL_PADDING_TWIPS,
                    bottom=const.CELL_VERTICAL_PADDING_TWIPS
                )

            # Apply width to the 4th column cell on this row if requested
            if fourth_col_width_inches is not None and len(df.columns) >= 4:
                cells[3].width = Inches(fourth_col_width_inches)

        # Extra safety: iterate all rows and enforce the width on col 4 (optional)
        if fourth_col_width_inches is not None and len(df.columns) >= 4:
            for r in table.rows:
                r.cells[3].width = Inches(fourth_col_width_inches)

    
    def set_cell_text(self, cell, text: str, bold: bool = False, size: Pt = Pt(10)):
            """
            Safely set a cell's text with styling without duplicating runs.
            - Clears cell content so the cell has a single paragraph.
            - Writes one run with the requested font attributes.
            """
            # Clear existing content to avoid duplicate runs
            cell.text = ""
            p = cell.paragraphs[0]

            # Use one run only
            if p.runs:
                r = p.runs[0]
                r.text = text
            else:
                r = p.add_run(text)

            r.font.bold = bold
            r.font.size = size

    def set_table_horizontal_borders(self, table, blank_only=False):
        """
        Keep only horizontal borders (top, bottom, insideH).
        Remove vertical borders (left, right, insideV).
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)

        def _border(tag, val='single', sz=8, color='auto', space='0'):
            el = tblBorders.find(qn(f'w:{tag}'))
            if el is None:
                el = OxmlElement(f'w:{tag}')
                tblBorders.append(el)
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), color)
            el.set(qn('w:space'), space)
            return el

        def _remove_border(tag):
            el = tblBorders.find(qn(f'w:{tag}'))
            if el is None:
                el = OxmlElement(f'w:{tag}')
                tblBorders.append(el)
            # 'nil' removes the border
            el.set(qn('w:val'), 'nil')

        if not blank_only:
            # Keep horizontal borders
            _border('top')
            _border('bottom')
            _border('insideH')
        else:
            # Remove horizontal borders
            _remove_border('top')
            _remove_border('bottom')
            _remove_border('insideH')

        # Remove vertical borders
        _remove_border('left')
        _remove_border('right')
        _remove_border('insideV')

    def set_cell_margins(self, cell, left=None, right=None, top=None, bottom=None):
        """
        Set table cell margins for a single cell using w:tcMar.
        All values are in twips (1 inch = 1440 twips).
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Get or create the <w:tcMar> element
        tcMar = tcPr.find(qn('w:tcMar'))
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)

        def _set_edge(edge_name, value):
            el = tcMar.find(qn(f'w:{edge_name}'))
            if el is None:
                el = OxmlElement(f'w:{edge_name}')
                tcMar.append(el)
            el.set(qn('w:w'), str(value))
            el.set(qn('w:type'), 'dxa')  # dxa = twips

        if left is not None:
            _set_edge('start', left)   # 'start' respects LTR/RTL; use 'left' if you prefer fixed side
        if right is not None:
            _set_edge('end', right)    # use 'right' for fixed side
        if top is not None:
            _set_edge('top', top)
        if bottom is not None:
            _set_edge('bottom', bottom)
            
    def set_cell_not_bold(self, cell):
        """
        Force all text in a cell to be non-bold by iterating paragraphs/runs.
        """
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = False

    def add_heading_generic(self, container, 
                            text: str, 
                            level: int = 2, 
                            alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            bold = True):
        """ Add a heading either to a Document or to a _Cell.
        For _Cell, we use a styled paragraph ('Heading X'). """
        
        if hasattr(container, "add_heading"):
            h = container.add_heading(text, level=level)
            h.alignment = alignment
            if not bold:
                for r in h.runs:
                    r.font.bold = False
        else:
            p = container.add_paragraph(text)
            p.style = f"Heading {level}"
            p.alignment = alignment
            if not bold:
                for r in p.runs:
                    r.font.bold = False
                
    def add_whitespace(
        self,
        container,
        space_before_pts: float = 6,
        space_after_pts: float = 6,
        add_divider: bool = False,
        divider_thickness_pts: float = 0.75,
        divider_indent_left_pts: float = 0,   # indent from left
        divider_indent_right_pts: float = 0,  # indent from right
        page_break: bool = False,
        keep_with_next: bool = False,
        keep_together: bool = False
    ):
        """
        Insert whitespace between components in a Document or a _Cell.

        Parameters
        ----------
        container : Document or _Cell
            Where the spacer should be inserted.

        space_before_pts : float
            Vertical space (points) before the spacer paragraph.

        space_after_pts : float
            Vertical space (points) after the spacer paragraph.

        add_divider : bool
            If True, draws a thin horizontal divider using a styled run.

        divider_thickness_pts : float
            Visual weight of the divider (simulated via font size and underline).

        divider_indent_left_pts, divider_indent_right_pts : float
            Pseudo-indentation by adding spaces before/after the divider run.
            Useful to shorten the divider relative to the container width.

        page_break : bool
            If True, inserts a page break *after* the spacer.

        keep_with_next : bool
            If True, prevents the spacer from being separated from the next block.

        keep_together : bool
            If True, tries to keep the spacer paragraph lines together.

        Returns
        -------
        Paragraph
            The spacer paragraph that was added (so you can further tweak it).
        """
        # Create a paragraph in the given container
        p = container.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Use a non-breaking space so Word preserves an "empty" paragraph
        r = p.add_run("\u00A0")
        r.font.size = Pt(1)  # tiny; visual height comes from paragraph spacing

        # Apply paragraph spacing
        pf = p.paragraph_format
        pf.space_before = Pt(space_before_pts)
        pf.space_after  = Pt(space_after_pts)
        pf.keep_with_next = keep_with_next
        pf.keep_together  = keep_together

        # Optional: light divider line
        if add_divider:
            # Simulate a thin rule using underline on a run of spaced characters
            # You can also replace with a series of '—' or '―' if preferred.
            left_pad  = " " * max(int(divider_indent_left_pts // 1), 0)
            right_pad = " " * max(int(divider_indent_right_pts // 1), 0)
            divider_text = f"{left_pad}" + ("_" * 40) + f"{right_pad}"  # 40 underscores ≈ mid-width
            r_div = p.add_run(divider_text)
            r_div.font.size = Pt(divider_thickness_pts)  # controls visual weight
            r_div.font.bold = False

        # Optional: page break *after* spacer
        if page_break:
            br = p.add_run()
            br.add_break(WD_BREAK.PAGE)

        return p